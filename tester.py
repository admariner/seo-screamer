from docx import Document
from docx.shared import RGBColor, Inches, Cm, Pt

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# doc = Document('word_templates/font_awesome.docx')
doc = Document()

obj_styles = doc.styles

obj_charstyle = obj_styles.add_style('FontAwesomeBrands', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(10)
obj_font.name = 'Font Awesome 5 Brands'

obj_charstyle = obj_styles.add_style('FontAwesomeLight', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(10)
obj_font.name = 'Font Awesome 5 Pro Light'

obj_charstyle = obj_styles.add_style('FontAwesomeRegular', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(10)
obj_font.name = 'Font Awesome 5 Pro Regular'

obj_charstyle = obj_styles.add_style('FontAwesomeSolid', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(10)
obj_font.name = 'Font Awesome 5 Pro Solid'


story = "Wordpress "
paragraph = doc.add_paragraph()
paragraph.add_run(story, style='FontAwesomeBrands')

story = "	alicorn f6b0 Light"
paragraph = doc.add_paragraph()
paragraph.add_run(story, style='FontAwesomeLight')

story = "	alicorn	f6b0 Regular"
paragraph = doc.add_paragraph()
paragraph.add_run(story, style='FontAwesomeRegular')


story = "	alicorn	f6b0 Solid"
paragraph = doc.add_paragraph()
paragraph.add_run(story, style='FontAwesomeSolid')


story = " desktop  wrench 	medal"
paragraph = doc.add_paragraph()
paragraph.add_run(story, style='FontAwesomeLight')


doc.save('test.docx')
