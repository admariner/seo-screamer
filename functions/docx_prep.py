from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE


class DocxPrep:
    def __init__(self, doc=None):
        if doc is None:
            return False

        self.doc = doc

    def set_document_styles(self):
        obj_styles = self.doc.styles

        style = self.doc.styles['Title']
        font = style.font
        font.name = 'Font Awesome 5 Pro Light'
        font.size = Pt(20)

        style = self.doc.styles['Heading 1']
        font = style.font
        font.name = 'Font Awesome 5 Pro Light'
        font.size = Pt(16)

        style = self.doc.styles['Heading 2']
        font = style.font
        font.name = 'Font Awesome 5 Pro Light'
        font.size = Pt(14)

        style = self.doc.styles['Heading 3']
        font = style.font
        font.name = 'Font Awesome 5 Pro Light'
        font.size = Pt(12)

        obj_charstyle = obj_styles.add_style('FontAwesomeBrands', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(13)
        obj_font.name = 'Font Awesome 5 Brands'

        obj_charstyle = obj_styles.add_style('FontAwesomeLight', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(13)
        obj_font.name = 'Font Awesome 5 Pro Light'

        obj_charstyle = obj_styles.add_style('FontAwesomeRegular', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(13)
        obj_font.name = 'Font Awesome 5 Pro Regular'

        obj_charstyle = obj_styles.add_style('FontAwesomeSolid', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(13)
        obj_font.name = 'Font Awesome 5 Pro Solid'

    def print_paragraph_styles(self):
        styles = self.doc.styles
        paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
        for style in paragraph_styles:
            print(style.name)

    def print_table_styles(self):
        styles = self.doc.styles
        paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.TABLE]
        for style in paragraph_styles:
            print(style.name)

    def print_list_styles(self):
        styles = self.doc.styles
        paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.LIST]
        for style in paragraph_styles:
            print(style.name)

    def empty_page(self, header=None):
        self.doc.add_page_break()
        self.doc.add_heading(header, level=1)
        self.doc.add_paragraph("Helaas hebben wij geen toegang tot uw website gegevens. Als u ons toegang geeft tot uw website "
                               "gegevens (lees rechten) dan kunnen wij nog meer informatie voor u ophalen zodat u uw "
                               "website nog beter kunt optimaliseren.")

        self.doc.add_paragraph("Wij kunnen u helpen om deze informatie voor ons inzichtelijk te maken.")


    def auto_cell(self, table):
        for r in table.rows:
            for c in r._tr.tc_lst:
                tcW = c.tcPr.tcW
                tcW.type = 'auto'
                tcW.w = 0
