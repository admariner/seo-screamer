import os
import sys
import logging
from docx import Document
from docx.shared import RGBColor, Inches, Cm, Pt

from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class toc:
    def __init__(self, doc=None):
        try:
            if doc is None:
                raise Exception("Doc is None")

            doc.add_page_break()
            doc.add_heading('Inhoudsopgave', 1)

            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')  # creates a new element
            fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
            instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # change 1-3 depending on heading levels you need

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            fldChar3 = OxmlElement('w:t')
            fldChar3.text = "Right-click to update field."
            fldChar2.append(fldChar3)

            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')

            r_element = run._r
            r_element.append(fldChar)
            r_element.append(instrText)
            r_element.append(fldChar2)
            r_element.append(fldChar4)
            p_element = paragraph._p
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            logging.warning(str(e) + " | " + str(exc_type) + " | " + str(fname) + " | " + str(exc_tb.tb_lineno))
            return None
