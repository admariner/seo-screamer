import sys
import os
import logging

from docx.shared import RGBColor, Pt, Cm
from docx.enum.table import WD_ALIGN_VERTICAL


class AlgemeenOverzicht:
    def __init__(self, doc=None, config=None, co_readydata=None, co_table_headers=None, descript_lang=None):
        try:
            if doc is None:
                raise Exception("Doc is None")

            crawl_overview_conf = descript_lang['crawl_overview']

            severity = config['severity']

            doc.add_page_break()
            doc.add_heading("Algemeen overzicht", level=1)

            par = doc.add_paragraph()
            par.add_run(
                "Achter iedere probleem melding vind u de oplossingsmoeilijkheid. Bij 1 ")
            run = par.add_run("")
            run.font.name = 'Font Awesome 5 Pro Light'
            par.add_run(
                " kunt u dit vrij eenvoudig zelf oplossen via de backend van uw site. Bij meer dan 3 ")
            run = par.add_run("")
            run.font.name = 'Font Awesome 5 Pro Light'
            par.add_run(" zal er een techneut aan te pas moeten komen.")

            if co_table_headers is None:
                raise Exception('Co Table Headers is None!')

            for h in co_table_headers:
                doc.add_paragraph()
                doc.add_heading(crawl_overview_conf[h]['name'], level=2)

                if crawl_overview_conf[h]['description'] != '':
                    doc.add_paragraph(crawl_overview_conf[h]['description'])

                if crawl_overview_conf[h]['recommendations'] is not None:
                    doc.add_paragraph()
                    doc.add_heading('Aanbevelingen', level=3)
                    for recomm in crawl_overview_conf[h]['recommendations']:
                        doc.add_paragraph(recomm, style='List Bullet')

                doc.add_paragraph()
                if crawl_overview_conf[h]['table_header'] == 1:
                    table = doc.add_table(
                        rows=1, cols=crawl_overview_conf[h]['table_cols'], style="Grid Table 4 Accent 5")
                    heading_cells = table.rows[0].cells
                    heading_cells[0].text = "Omschrijving"
                    heading_cells[1].text = "Aantal"
                    if crawl_overview_conf[h]['table_cols'] > 2:
                        heading_cells[2].text = "Uitkomst"
                    if crawl_overview_conf[h]['table_cols'] > 3:
                        heading_cells[3].text = "Moeilijkheid"

                else:
                    table = doc.add_table(
                        rows=0, cols=crawl_overview_conf[h]['table_cols'])

                for r in co_readydata[h]:
                    row_cells = table.add_row().cells
                    row_cells[0].text = r[0]
                    row_cells[1].text = r[1]

                    if crawl_overview_conf[h]['table_cols'] > 2:
                        paragraph = row_cells[2].paragraphs[0]
                        if r[2] == '':
                            paragraph.add_run("")
                        elif r[2] == 0:
                            runner = paragraph.add_run(severity['low']['icon'])
                            runner.font.name = 'Font Awesome 5 Pro Light'
                            color = severity['low']['color']
                            runner.font.color.rgb = RGBColor(
                                color[0], color[1], color[2])
                            runner.font.size = Pt(13)
                            runner.bold = True
                        elif r[2] == 1 or r[2] == 2:
                            runner = paragraph.add_run(
                                severity['medium']['icon'])
                            runner.font.name = 'Font Awesome 5 Pro Light'
                            color = severity['medium']['color']
                            runner.font.color.rgb = RGBColor(
                                color[0], color[1], color[2])
                            runner.font.size = Pt(13)
                            runner.bold = True
                        elif r[2] > 2:
                            runner = paragraph.add_run(
                                severity['high']['icon'])
                            runner.font.name = 'Font Awesome 5 Pro Light'
                            color = severity['high']['color']
                            runner.font.color.rgb = RGBColor(
                                color[0], color[1], color[2])
                            runner.font.size = Pt(13)
                            runner.bold = True

                    if crawl_overview_conf[h]['table_cols'] > 3:
                        try:
                            paragraph = row_cells[3].paragraphs[0]
                            runner = paragraph.add_run(r[3])
                            nr = round(len(list(r[3]))/2)
                            if nr <= 2:
                                color = severity['low']['color']
                            elif 2 < nr <= 3:
                                color = severity['medium']['color']
                            elif nr > 3:
                                color = severity['high']['color']

                            runner.font.name = 'Font Awesome 5 Pro Light'
                            runner.font.color.rgb = RGBColor(
                                color[0], color[1], color[2])

                        except IndexError as e:
                            raise Exception(
                                "We found an error in {} : {}".format(r[0], e))

                self.set_col_widths(
                    table, crawl_overview_conf[h]['table_cols'])
                self.set_row_heights(table, 0.7)
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            logging.warning(str(e) + " | " + str(exc_type) +
                            " | " + str(fname) + " | " + str(exc_tb.tb_lineno))
            return None

    def set_col_widths(self, table, cells=4):
        widths = (Cm(8), Cm(2), Cm(2), Cm(4))
        if cells == 2:
            widths = (Cm(10), Cm(6))
        if cells == 3:
            widths = (Cm(10), Cm(3), Cm(3))
        if cells == 4:
            widths = (Cm(8), Cm(2), Cm(2), Cm(4))

        i = 0
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
                row.cells[idx].height = Cm(1.5)
                table.cell(
                    i, idx).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            i += 1

    def set_row_heights(self, table, height=0.7):
        for row in table.rows:
            row.height = Cm(height)
