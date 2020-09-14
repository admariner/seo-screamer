import sys
import os
import logging

from docx.shared import RGBColor, Inches, Cm, Pt

# from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.text import WD_COLOR_INDEX
#
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .page_speed import PageSpeed


class DocxPageSpeed:

    def __init__(self, doc=None, config=None, url=None, domain=None, strategy='mobile'):
        try:
            if config is None:
                raise Exception("Config is None")
            if url is None:
                raise Exception("Url is None")
            if domain is None:
                raise Exception("Domain is None")
            if doc is None:
                raise Exception("Doc is None")

            self.doc = doc
            self.ps = None
            self.ps_strategy = strategy
            self.ps_config = config

            self.ps_api = self.ps_config['google_page_speed_api']
            self.ps_url = url
            self.ps_domain = domain

            self.get_page_speed_data()
            self.create_doc_page()
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            logging.warning(str(e) + " | " + str(exc_type) +
                            " | " + str(fname) + " | " + str(exc_tb.tb_lineno))
            return None

    def create_doc_page(self):
        self.doc.add_page_break()

        icon = self.ps_config['icons']['desktop']
        if self.ps_strategy == 'mobile':
            icon = self.ps_config['icons']['mobile']

        self.doc.add_heading("{} {} Test".format(
            icon, self.ps_strategy.capitalize()), level=1)

        self.doc.add_paragraph(
            'Onderstaande tests zijn gedaan op basis van de homepage.')

        table = self.doc.add_table(rows=1, cols=2)
        for v in self.ps.category:
            row_cells = table.add_row().cells
            row_cells[0].text = self.ps.categorien[v]
            paragraph = row_cells[1].paragraphs[0]
            paragraph.add_run("{}".format(self.ps.ready_data['{}_{}_score'.format(self.ps_strategy, v)])).font.color.rgb = \
                self.colorsPercent(
                    self.ps.ready_data['{}_{}_score'.format(self.ps_strategy, v)])

        self.doc.add_paragraph()

        table = self.doc.add_table(rows=1, cols=4)
        table.allow_autofit = True
        row_cells = table.add_row().cells
        row_cells[0].text = "Schaal"

        paragraph = row_cells[1].paragraphs[0]
        runner = paragraph.add_run("●")
        runner.font.color.rgb = RGBColor(0, 131, 48)
        runner.font.size = Pt(16)
        paragraph.add_run(" 90-100")

        paragraph = row_cells[2].paragraphs[0]
        runner = paragraph.add_run("●")
        runner.font.color.rgb = RGBColor(255, 162, 0)
        runner.font.size = Pt(16)
        paragraph.add_run(" 50-89")

        paragraph = row_cells[3].paragraphs[0]
        runner = paragraph.add_run("●")
        runner.font.color.rgb = RGBColor(161, 33, 1)
        runner.font.size = Pt(16)
        paragraph.add_run(" 0-49")

        row_cells = table.add_row().cells
        row_cells[0].text = ""
        row_cells[1].text = "(snel)"
        row_cells[2].text = "(gemiddeld)"
        row_cells[3].text = "(langzaam)"

        self.doc.add_paragraph()

        self.doc.add_paragraph("Analyse van de pagina via een geëmuleerd netwerk. Waarden worden geschat en "
                               "kunnen variëren.").italic = True

        self.doc.add_paragraph()
        icon = self.ps_config['icons']['stopwatch']
        self.doc.add_heading("{} Metrieken".format(icon), level=2)

        table = self.doc.add_table(rows=1, cols=4)

        i = 0
        for a in self.ps.audit_list_performance:
            if i == 0:
                row_cells = table.add_row().cells

            row_cells[i].text = self.ps.ready_data['{}_{}_title'.format(
                self.ps_strategy, a)]
            i += 1
            paragraph = row_cells[i].paragraphs[0]
            paragraph.add_run("{}".format(self.ps.ready_data['{}_{}_displayValue'.format(self.ps_strategy, a)])).font.\
                color.rgb = self.colorsPercent(
                    self.ps.ready_data['{}_{}_score'.format(self.ps_strategy, a)])
            i += 1
            if i == 4:
                i = 0

        self.auto_cell(table)

        self.doc.add_paragraph()
        icon = self.ps_config['icons']['ballot']
        self.doc.add_heading("{} Aanbevelingen".format(icon), level=2)
        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "Deze optimalisaties kunnen het laden van je pagina versnellen.")
        self.doc.add_paragraph()

        i = 0
        table = self.doc.add_table(
            rows=1, cols=2, style="Table Grid")  # , style='Table Grid' ,Seo_pagespeed
        heading_cells = table.rows[0].cells
        heading_cells[0].text = "Aanbeveling"
        heading_cells[1].text = "Geschatte besparing"

        self.shade_cells([table.cell(0, 0), table.cell(0, 1)], "#99badd")

        for r in self.ps.recommendations:
            i += 2
            cells = table.add_row().cells
            cells[0].text = r['title']
            cells[1].text = str(r['displayValue'])

            cells = table.add_row().cells
            cells[0].text = r['descr']
            cells[1].text = ""
            a = table.cell(i, 0)
            b = table.cell(i, 1)
            a.merge(b)

        self.auto_cell(table)

        self.doc.add_paragraph()
        icon = self.ps_config['icons']['stethoscope']
        self.doc.add_heading("{} Diagnostische gegevens".format(icon), level=2)
        self.doc.add_paragraph()
        self.doc.add_paragraph(
            "Meer informatie over de prestaties van je app.")
        self.doc.add_paragraph()

        table = self.doc.add_table(
            rows=1, cols=2, style="Table Grid")  # Seo_pagespeed
        heading_cells = table.rows[0].cells
        heading_cells[0].text = "Aanbeveling"
        heading_cells[1].text = "Geschatte besparing"

        i = 0
        for d in self.ps.diagnostic_data:
            i += 2
            cells = table.add_row().cells
            cells[0].text = d['title']
            cells[1].text = "{}".format(d['displayValue'])

            cells = table.add_row().cells
            cells[0].text = d['descr']
            cells[1].text = ""
            a = table.cell(i, 0)
            b = table.cell(i, 1)
            a.merge(b)

        self.auto_cell(table)

    def get_page_speed_data(self):
        self.ps = PageSpeed(self.ps_api, self.ps_url,
                            self.ps_domain, self.ps_strategy)
        self.ps.get_audits()

    def auto_cell(self, table):
        for r in table.rows:
            for c in r._tr.tc_lst:
                tcW = c.tcPr.tcW
                tcW.type = 'auto'
                tcW.w = 0

    def shade_cells(self, cells, shade):
        for cell in cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcVAlign = OxmlElement("w:shd")
            tcVAlign.set(qn("w:fill"), shade)
            tcPr.append(tcVAlign)

    def colorsPercent(self, number=0):
        # https://www.colorspire.com/rgb-color-wheel/
        if number < 50:
            return RGBColor(161, 33, 1)
        elif number >= 50 and number < 89:
            return RGBColor(255, 162, 0)
        elif number >= 89:
            return RGBColor(0, 131, 48)
