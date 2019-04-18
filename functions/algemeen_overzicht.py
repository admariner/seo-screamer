from docx.shared import RGBColor, Pt, Cm
from docx.enum.table import WD_ALIGN_VERTICAL

class AlgemeenOverzicht:
    def __init__(self, doc=None, config=None, co_readydata=None, co_table_headers=None):
        if doc is None:
            return False

        crawl_overview_conf = config['crawl_overview']
        severity = config['severity']

        doc.add_page_break()
        doc.add_heading("Algemeen overzicht", level=1)

        for h in co_table_headers:
            doc.add_paragraph()
            doc.add_heading(h, level=2)

            if crawl_overview_conf[h]['description'] != '':
                doc.add_paragraph(crawl_overview_conf[h]['description'])

            if crawl_overview_conf[h]['recommendations'] is not None:
                doc.add_paragraph()
                doc.add_heading('Aanbevelingen', level=3)
                for recomm in crawl_overview_conf[h]['recommendations']:
                    doc.add_paragraph(recomm, style='List Bullet')

            table = doc.add_table(rows=1, cols=3)
            for r in co_readydata[h]:
                row_cells = table.add_row().cells
                row_cells[0].text = r[0]
                row_cells[1].text = r[1]
                paragraph = row_cells[2].paragraphs[0]
                if r[2] == '':
                    paragraph.add_run("")
                elif r[2] == 0:
                    runner = paragraph.add_run(severity['low']['icon'])
                    runner.font.name = 'Font Awesome 5 Pro'
                    color = severity['low']['color']
                    runner.font.color.rgb = RGBColor(color[0], color[1], color[2])
                    runner.font.size = Pt(13)
                    runner.bold = True
                elif r[2] == 1 or r[2] == 2:
                    runner = paragraph.add_run(severity['medium']['icon'])
                    runner.font.name = 'Font Awesome 5 Pro'
                    color = severity['medium']['color']
                    runner.font.color.rgb = RGBColor(color[0], color[1], color[2])
                    runner.font.size = Pt(13)
                    runner.bold = True
                elif r[2] > 2:
                    runner = paragraph.add_run(severity['high']['icon'])
                    runner.font.name = 'Font Awesome 5 Pro'
                    color = severity['high']['color']
                    runner.font.color.rgb = RGBColor(color[0], color[1], color[2])
                    runner.font.size = Pt(13)
                    runner.bold = True

            for r in table.rows:
                r.height = Cm(0.7)
                for cell in r.cells:
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW = tcPr.get_or_add_tcW()
                    tcW.type = 'auto'
