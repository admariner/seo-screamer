import os
import sys
import logging

from docx import Document
from docx.shared import RGBColor, Inches, Cm, Pt

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from functions.csv_files import ParceCSV


class CSV2Docx:
    def __init__(self, config=None, doc=None, frog_data_folder=None, export_tabs=None, bulk_exports=None):
        try:
            if config is None:
                raise Exception("Config is None")
            if doc is None:
                raise Exception("Doc is None")
            if frog_data_folder is None:
                raise Exception("frog_data_folder is None")
            if export_tabs is None:
                raise Exception("Export_tabs is None")
            if bulk_exports is None:
                raise Exception("Bulk exports is None")

            export_tabs = [i for i in export_tabs['export_tabs']
                           if not (i['active'] == 0)]
            bulk_exports = [
                i for i in bulk_exports['bulk_exports'] if not (i['active'] == 0)]

            crawl_files = export_tabs + bulk_exports

            doc.add_heading('Crawl overzichten', 0)
            doc.add_paragraph("Hier vind u alle informatie behorende bij het algemene overzicht. "
                              "Deze informatie geeft u meer inzicht in wat u inhoudelijk aan uw "
                              "pagina's dient te wijzigen om betere SEO resultaten te krijgen.")

            for cf in crawl_files:

                doc.add_paragraph()
                file = os.path.join(frog_data_folder, cf['file'])
                c = ParceCSV(file)

                for key, val in c.ready_data.items():
                    if len(val['data']) == 0:
                        continue

                    if cf['id'] == "":
                        doc.add_heading(key, 1)
                    else:
                        doc.add_heading(cf['id'], 1)

                    try:
                        doc.add_paragraph(cf['description'])
                    except Exception:
                        pass

                    try:
                        kolommen = len(cf['columns'])
                    except Exception as e:
                        print(e)
                        sys.exit()

                    table = doc.add_table(
                        rows=1, cols=kolommen, style="Grid Table 4 Accent 5")
                    hdr_cells = table.rows[0].cells
                    self.set_repeat_table_header(table.rows[0])
                    i = 0
                    for h in cf['columns']:
                        cell = hdr_cells[i]
                        paragraph = cell.paragraphs[0]
                        runner = paragraph.add_run(h)
                        runner.font.size = Pt(9)
                        i += 1

                    for d in val['data']:
                        row_cells = table.add_row().cells
                        i = 0
                        for h in cf['columns']:
                            cell = row_cells[i]
                            paragraph = cell.paragraphs[0]
                            runner = paragraph.add_run(str(d[h]))
                            runner.font.size = Pt(9)
                            i += 1

                    self.auto_cell(table)
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            logging.warning(str(e) + " | " + str(exc_type) +
                            " | " + str(fname) + " | " + str(exc_tb.tb_lineno))
            return None

    def auto_cell(self, table):
        for r in table.rows:
            for c in r._tr.tc_lst:
                tcW = c.tcPr.tcW
                tcW.type = 'auto'
                tcW.w = 0

    def set_repeat_table_header(self, row):
        """ set repeat table row on every new page
        """
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row
