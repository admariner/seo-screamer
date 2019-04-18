# pip install python-docx

import os
from datetime import date

from docx import Document
from docx.shared import RGBColor, Inches, Cm, Pt

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from functions.readConfig import readConfig
from functions.crawl_overview import CrawlOverview
from functions.csv_files import ParceCSV
from functions.page_speed import PageSpeed
from functions.inleiding import inleiding
from functions.eerste_pagina import EerstePagina
from functions.algemeen_overzicht import AlgemeenOverzicht

# import locale
# locale.setlocale(locale.LC_TIME, "nl_NL.utf8")

class CreateWord:
    def __init__(self, config=None):
        self.datum = '{:%d-%b-%Y}'.format(date.today())

        try:
            self.url = config['url']
            self.domain = config['domain']
            self.tempate_file = config['word_template']
            self.search_console_url = config['search_console_url']
        except KeyError as e:
            return "Config variables not found! {}".format(e)

        self.frog_files = {}
        self.frog_data_folder = None
        self.document_fields = None
        self.doc = None
        self.data = {}
        self.co_table_headers = None
        self.co_readydata = None
        self.template = None

        cf = readConfig()
        config = cf.config
        self.crawl_files = [i for i in config['crawl_files'] if not (i['active'] == 0)]
        # self.crawl_files = crawl_files = [i for i in config['crawl_files'] if not (i['active'] == 0)]
        self.config = config
        self.ps_api = config['google_page_speed_api']

        self.ps_category = []
        self.ps_categorien = {}
        self.audit_list_performance = []

        self.word_output_file = self.set_doc_params()
        self.frog_data_folder = self.get_frog_folder()
        self.get_frog_files()


        self.get_crawl_overview_data()
        self.ps_all_data = {}
        self.ps_data = {'mobile': self.get_page_speed_data('mobile'), 'desktop': self.get_page_speed_data('desktop')}
        # self.crawl_data = self.get_crawl_data()
        
        self.open_document()
        self.set_document_styles()
        self.merge_document()
        self.write_document()
    
    def set_doc_params(self):
        dir_path = os.path.dirname(os.path.realpath(__file__))

        if self.tempate_file is not None:
            template_folder = os.path.join(dir_path, "word_templates")
            self.template = os.path.join(template_folder, self.tempate_file)
        
        output_folder = os.path.join(dir_path, "word_output")
        output_name = "{}-{}.docx".format(self.datum, self.domain)
        return os.path.join(output_folder, output_name)
    
    def get_frog_folder(self):
        dir_path = os.path.dirname(os.path.realpath(__file__))
        data_folder = os.path.join(dir_path, "data")
        return os.path.join(data_folder, self.domain)

    def get_frog_files(self):
        for r, d, f in os.walk(self.frog_data_folder):
            for file in f:
                if '.csv' in file:
                    self.frog_files[file] = os.path.join(r, file)

    def get_crawl_overview_data(self):
        try:
            co = CrawlOverview(self.frog_files['crawl_overview.csv'])
            self.co_table_headers = co.table_headers
            self.co_readydata = co.readydata
        except KeyError:
            print('error')
            return {}

    def get_page_speed_data(self, strategy='mobile'):
        ps = PageSpeed(self.ps_api, self.url, self.domain, strategy)
        self.ps_category = ps.category
        self.ps_categorien = ps.categorien
        self.audit_list_performance = ps.audit_list_performance
        self.ps_all_data = ps.ps_data
        return ps.ready_data

    def open_document(self):
        self.doc = Document(self.template)

    def set_document_styles(self):
        obj_styles = self.doc.styles

        obj_charstyle = obj_styles.add_style('FontAwesomeBrands', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(10)
        obj_font.name = 'Font Awesome 5 Brands'

        obj_charstyle = obj_styles.add_style('FontAwesomeLight', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(10)
        obj_font.name = 'Font Awesome 5 Pro Light'

        obj_charstyle = obj_styles.add_style('FontAwesomeRegular', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(10)
        obj_font.name = 'Font Awesome 5 Pro Regular'

        obj_charstyle = obj_styles.add_style('FontAwesomeSolid', WD_STYLE_TYPE.CHARACTER)
        obj_font = obj_charstyle.font
        obj_font.size = Pt(10)
        obj_font.name = 'Font Awesome 5 Pro Solid'

    def print_document_styles(self):
        styles = self.doc.styles
        paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
        for style in paragraph_styles:
            print(style.name)    

    def merge_document(self):
        sc = 0
        EerstePagina(self.doc, sc, self.domain)
        inleiding(self.doc, self.domain)
        AlgemeenOverzicht(self.doc, self.config, self.co_readydata, self.co_table_headers)

        self.page_speed('mobile')
        self.page_speed('desktop')
        
        self.change_orientation()

        self.doc.add_heading('Crawl overzichten', 0)
        self.doc.add_paragraph("Hier vind u alle informatie behorende bij het algemene overzicht. "
                                "Deze informatie geeft u meer inzicht in wat u inhoudelijk aan uw "
                                "pagina's dient te wijzigen om betere SEO resultaten te krijgen")
                                
        for cf in self.crawl_files:
            
            self.doc.add_paragraph()
            
            file = os.path.join(self.frog_data_folder, cf['file'])
            c = ParceCSV(file)
        
            for key, val in c.ready_data.items():
                if len(val['data']) == 0:
                    continue
                
                if cf['name'] == "":
                    self.doc.add_heading(key, 1)
                else:
                    self.doc.add_heading(cf['name'], 1)
                    
                self.doc.add_paragraph(cf['description'])
                                
                kolommen = len(cf['columns'])
                table = self.doc.add_table(rows=1, cols=kolommen)
                hdr_cells = table.rows[0].cells
                self.set_repeat_table_header(table.rows[0])
                i = 0
                for h in cf['columns']:
                    cell = hdr_cells[i]
                    paragraph = cell.paragraphs[0]
                    runner = paragraph.add_run(h)
                    runner.font.size = Pt(9)
                    i += 1
                
                for d in val['data']:
                    row_cells = table.add_row().cells
                    i = 0
                    for h in cf['columns']:
                        cell = row_cells[i]
                        paragraph = cell.paragraphs[0]
                        runner = paragraph.add_run(str(d[h]))
                        runner.font.size = Pt(9)
                        i += 1
            
                for r in table.rows:
                    for c in r._tr.tc_lst:
                        tcW = c.tcPr.tcW
                        tcW.type = 'auto'
                        tcW.w = 0
        # self.legenda()

    def page_speed(self, strategy='mobile'):
        self.doc.add_page_break()
        self.doc.add_heading("{} Test".format(strategy.capitalize()), level=1)

        self.doc.add_paragraph('Onderstaande tests zijn gedaan op basis van de homepage.')

        table = self.doc.add_table(rows=1, cols=2)
        for v in self.ps_category:
            row_cells = table.add_row().cells
            row_cells[0].text = self.ps_categorien[v]
            paragraph = row_cells[1].paragraphs[0]
            paragraph.add_run("{}".format(self.ps_data[strategy]['{}_{}_score'.format(strategy, v)])).font.color.rgb = \
                self.colorsPercent(self.ps_data[strategy]['{}_{}_score'.format(strategy, v)])

        self.doc.add_paragraph()

        table = self.doc.add_table(rows=1, cols=4)
        table.allow_autofit = True
        row_cells = table.add_row().cells
        row_cells[0].text = "Schaal"

        paragraph = row_cells[1].paragraphs[0]
        runner = paragraph.add_run("●")
        runner.font.color.rgb = RGBColor(0, 131, 48)
        runner.font.size = Pt(16)
        paragraph.add_run(" 90-100")

        paragraph = row_cells[2].paragraphs[0]
        runner = paragraph.add_run("●")
        runner.font.color.rgb = RGBColor(255, 162, 0)
        runner.font.size = Pt(16)
        paragraph.add_run(" 50-89")

        paragraph = row_cells[3].paragraphs[0]
        runner = paragraph.add_run("●")
        runner.font.color.rgb = RGBColor(161, 33, 1)
        runner.font.size = Pt(16)
        paragraph.add_run(" 0-49")

        row_cells = table.add_row().cells
        row_cells[0].text = ""
        row_cells[1].text = "(snel)"
        row_cells[2].text = "(gemiddeld)"
        row_cells[3].text = "(langzaam)"

        self.doc.add_paragraph()

        self.doc.add_paragraph("Analyse van de pagina via een geëmuleerd netwerk. Waarden worden geschat en "
                               "kunnen variëren.").italic = True

        self.doc.add_paragraph()
        self.doc.add_heading("Metrieken", level=2)

        table = self.doc.add_table(rows=1, cols=4)

        i = 0
        for a in self.audit_list_performance:
            if i == 0:
                row_cells = table.add_row().cells

            row_cells[i].text = self.ps_data[strategy]['{}_{}_title'.format(strategy, a)]
            i += 1
            paragraph = row_cells[i].paragraphs[0]
            paragraph.add_run("{}".format(self.ps_data[strategy]['{}_{}_displayValue'.format(strategy, a)])).font.\
                color.rgb = self.colorsPercent(self.ps_data[strategy]['{}_{}_score'.format(strategy, a)])
            i += 1
            if i == 4:
                i = 0

        for r in table.rows:
            for c in r._tr.tc_lst:
                tcW = c.tcPr.tcW
                tcW.type = 'auto'
                tcW.w = 0
       
    def legenda(self):
        self.change_orientation()
        self.doc.add_heading("Legenda & Uitleg", level=2)
        self.doc.add_paragraph()

        for key, val in self.ps_all_data['lighthouseResult']['audits'].items():
            self.doc.add_heading(val['title'], level=3)
            self.doc.add_paragraph(val['description'])
            self.doc.add_paragraph()       
            
    def change_orientation(self):
        current_section = self.doc.sections[-1]
        new_width, new_height = current_section.page_height, current_section.page_width
        new_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = new_width
        new_section.page_height = new_height
        return new_section            
            
    def colorsPercent(self, number=0):
        #https://www.colorspire.com/rgb-color-wheel/
        if number < 50:
            return RGBColor(161, 33, 1)
        elif number >= 50 and number < 89:
            return RGBColor(255, 162, 0)
        elif number >= 89:  
            return RGBColor(0, 131, 48)

    def set_repeat_table_header(self, row):
        """ set repeat table row on every new page
        """
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    def write_document(self):
        self.doc.save(self.word_output_file)


if __name__ == '__main__':
    dir_path = os.path.dirname(os.path.realpath(__file__))
    profiles = os.path.join(dir_path, "data")

    #
    # for domain in os.listdir(folder):
    #     domain_folder = os.path.join(folder, domain)
    #     if os.path.isdir(domain_folder) or os.path.islink(domain_folder):
    #         url = "https://{}".format(domain)
    #         t = None  # test_doc.docx   seo_doc.docx
    #         d = Domain(url)
    #         m = CreateWord(url, domain, t)

    profile = os.path.join(profiles, "oesterbaron.nl")
    config_file = os.path.join(profile, "config.yml")
    c = readConfig(config_file)
    m = CreateWord(c.config)
