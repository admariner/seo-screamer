# pip install python-docx

import os
import sys
import logging
import json

from time import time
from datetime import date, datetime
from docx import Document
from docx.shared import RGBColor, Inches, Cm, Pt

from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.text import WD_COLOR_INDEX
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from functions.docx_prep import DocxPrep
from functions.readConfig import readConfig
from functions.crawl_overview import CrawlOverview
from functions.google_search_console import GoogleSearchConsole
from functions.docx_inleiding import inleiding
from functions.docx_eerste_pagina import EerstePagina
from functions.docx_algemeen_overzicht import AlgemeenOverzicht
from functions.docx_toc import toc
from functions.docx_pagespeed import DocxPageSpeed
from functions.docx_csv_files import CSV2Docx
from functions.screaming import Screaming
from functions.docx_google_search_console import DocxGoogleSearch

# import locale
# locale.setlocale(locale.LC_TIME, "nl_NL.utf8")
now = datetime.now()
today = "{}-{}-{}".format(now.year, now.month, now.day)
sf = os.path.dirname(os.path.realpath(__file__))
folder = os.path.join(sf, 'logging')
log_file = os.path.join(folder, "{}.log".format(today))

if os.environ['HOME'] == '/Users/theovandersluijs':
    logging.basicConfig(format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                        level=logging.DEBUG)
else:
    os.makedirs(folder, exist_ok=True)
    logging.basicConfig(filename=log_file, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
                        level=logging.CRITICAL)


class CreateWord:
    def __init__(self, config=None, export_tabs=None, bulk_export=None,
                 crawl_yml: dict = None, descript_lang: dict = None):
        try:
            self.datum = '{:%d-%b-%Y}'.format(date.today())
            self.maand_jaar = '{:%B-%Y}'.format(date.today())

            try:
                self.url = config['url']
                self.domain = config['domain']
                self.template_file = config['word_template']
                self.search_console_url = config['search_console_url']
            except KeyError as e:
                raise Exception("Error getting config variables: {}".format(e))

            self.start_path = os.path.dirname(os.path.realpath(__file__))
            self.domain_folder = self.set_data_domain_path()
            self.frog_data_folder = self.get_frog_folder()
            self.graph_data_folder = self.get_graph_folder()
            self.config_folder = self.get_config_folder()
            self.export_tabs = export_tabs
            self.bulk_export = bulk_export
            self.crawl_yml = crawl_yml
            self.descript_lang = descript_lang

            self.ps = None
            self.frog_files = {}
            self.document_fields = None
            self.doc = None
            self.dp = None
            self.data = {}
            self.co_table_headers = None
            self.co_readydata = None
            self.template = None
            self.dimensions_data = None

            cf = readConfig()
            self.config = cf.config
            self.ps_api = self.config['google_page_speed_api']

            self.template = self.get_doc_template()
            self.word_output_file = self.get_doc_word_output()

            self.get_frog_files()
            self.create_crawl_data()
            self.get_crawl_overview_data()
            self.get_google_search_console_data()

            self.ps_all_data = {}

            # self.crawl_data = self.get_crawl_data()

            self.prep_document()
            self.merge_document()
            self.write_document()
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            logging.warning(str(e) + " | " + str(exc_type) +
                            " | " + str(fname) + " | " + str(exc_tb.tb_lineno))
            return

    def get_doc_template(self):
        if self.template_file is not None:
            return os.path.join(self.start_path, "word_templates", self.template_file)

    def get_doc_word_output(self):
        return os.path.join(self.start_path, "word_output", "{}-{}.docx".format(self.maand_jaar, self.domain))

    def get_config_folder(self):
        return os.path.join(self.start_path, 'config')

    def set_data_domain_path(self):
        return os.path.join(self.start_path, "data", self.domain)

    def get_graph_folder(self):
        return os.path.join(self.domain_folder, 'graphs')

    def get_frog_folder(self):
        return os.path.join(self.domain_folder, 'crawl')

    def get_frog_files(self):
        try:
            for r, d, f in os.walk(self.frog_data_folder):
                for file in f:
                    if '.csv' in file:
                        self.frog_files[file] = os.path.join(r, file)

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            logging.warning(str(e) + " | " + str(exc_type) +
                            " | " + str(fname) + " | " + str(exc_tb.tb_lineno))

    def create_crawl_data(self):
        try:
            t = 0
            try:
                if os.path.isfile(self.frog_files['crawl_overview.csv']):
                    t = os.path.getmtime(self.frog_files['crawl_overview.csv'])
            except KeyError:
                t = 0

            if int(t) + 86400 <= int(time()):
                print('Creating Screaming Frog crawl data')
                export_tabs = []
                bulk_export = []

                for e in self.export_tabs['export_tabs']:
                    if e['active'] == 1:
                        export_tabs.append(e['id'])

                for b in self.bulk_export['bulk_exports']:
                    if b['active'] == 1:
                        bulk_export.append(b['id'])

                s = Screaming(self.domain_folder, self.url,
                              self.search_console_url, export_tabs, bulk_export)
                s.run_screamer()
            else:
                print('Using current Screaming Frog crawl data')
        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            logging.critical(str(e) + " | " + str(exc_type) +
                             " | " + str(fname) + " | " + str(exc_tb.tb_lineno))
            sys.exit()

    def get_crawl_overview_data(self):
        try:
            try:
                co = CrawlOverview(
                    crawl_file=self.frog_files['crawl_overview.csv'], crawl_yml=self.crawl_yml)
            except KeyError as e:
                raise Exception("Error getting frog files key: {}".format(e))
            self.co_table_headers = co.table_headers
            self.co_readydata = co.readydata

        except Exception as e:
            exc_type, exc_obj, exc_tb = sys.exc_info()
            fname = os.path.split(exc_tb.tb_frame.f_code.co_filename)[1]
            logging.critical(str(e) + " | " + str(exc_type) +
                             " | " + str(fname) + " | " + str(exc_tb.tb_lineno))
            return {}

    def get_google_search_console_data(self):
        self.dimensions_data = ""

        if self.search_console_url is not None and self.search_console_url != "":
            dimension = ['device', 'query', 'country', 'page']
            g = GoogleSearchConsole(self.search_console_url, self.domain, date.today(),
                                    -30, dimension, self.config_folder)
            self.dimensions_data = g.dimension_data

    def prep_document(self):
        self.doc = Document(self.template)
        self.dp = DocxPrep(self.doc)
        self.dp.set_document_styles()

    def merge_document(self):
        sc = 0
        EerstePagina(self.doc, sc, self.domain)
        inleiding(self.doc, self.domain)
        toc(self.doc)

        # crawl overzicht
        AlgemeenOverzicht(doc=self.doc, config=self.config,
                          co_readydata=self.co_readydata,
                          co_table_headers=self.co_table_headers,
                          descript_lang=self.descript_lang)

        # Doc Pagespeed overzicht
        DocxPageSpeed(self.doc, self.config, self.url, self.domain, 'mobile')
        DocxPageSpeed(self.doc, self.config, self.url, self.domain, 'desktop')

        if self.dimensions_data != "":  # only use when we have data!
            DocxGoogleSearch(self.doc, self.dimensions_data,
                             self.graph_data_folder)
        else:
            self.dp.empty_page('Google Search Console')

        self.change_orientation()
        CSV2Docx(self.config, self.doc, self.frog_data_folder,
                 self.export_tabs, self.bulk_export)

        # self.legenda()

    def legenda(self):
        self.change_orientation()
        self.doc.add_heading("Legenda & Uitleg", level=2)
        self.doc.add_paragraph()

        for key, val in self.ps_all_data['lighthouseResult']['audits'].items():
            self.doc.add_heading(val['title'], level=3)
            self.doc.add_paragraph(val['description'])
            self.doc.add_paragraph()

        self.doc.add_heading("Google Zoeken", level=3)
        self.doc.add_heading("Klikken (CLicks)", level=4)
        self.doc.add_paragraph("Het aantal klikken (clicks) afkomstig van een Google zoekresultaat waarmee de "
                               "gebruiker/bezoeker op je site is terechtgekomen.")

        self.doc.add_heading("Vertoningen", level=4)
        self.doc.add_paragraph(
            "Hoeveel links naar je site een gebruiker zag op de pagina met Google zoekresultaten van "
            "Google Zoeken. Vertoningen worden geteld wanneer de gebruiker die pagina met "
            "resultaten bezoekt, zelfs wanneer de gebruiker niet naar het resultaat is gescrold. "
            "Als een gebruiker echter alleen pagina 1 bekijkt en het resultaat op pagina 2 staat, "
            "telt die vertoning niet mee.")

        self.doc.add_heading("Positie", level=4)
        self.doc.add_paragraph("De gemiddelde positie van het hoogste resultaat van de site. "
                               "Als je site bijvoorbeeld drie resultaten heeft op posities 2, 4 en 6, wordt de "
                               "positie gerapporteerd als 2. Belangrijk is dat deze waarde zo laag mogelijk is, "
                               "hoe lager hoe hoger in de zoek resultaten.")

    def change_orientation(self):
        current_section = self.doc.sections[-1]
        new_width, new_height = current_section.page_height, current_section.page_width
        new_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = new_width
        new_section.page_height = new_height
        return new_section

    def set_repeat_table_header(self, row):
        """ set repeat table row on every new page
        """
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    def shade_cells(self, cells, shade):
        for cell in cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcVAlign = OxmlElement("w:shd")
            tcVAlign.set(qn("w:fill"), shade)
            tcPr.append(tcVAlign)

    def write_document(self):
        self.doc.save(self.word_output_file)


if __name__ == '__main__':
    dir_path = os.path.dirname(os.path.realpath(__file__))
    profiles = os.path.join(dir_path, "data")

    word_output = os.path.join(dir_path, "word_output")
    os.makedirs(word_output, exist_ok=True)

    bulk_export = None
    bulk_export_file = os.path.join(dir_path, 'config', 'bulk_exports.yml')
    if os.path.isfile(bulk_export_file):
        bulk_export = readConfig(bulk_export_file)

    crawl_yml = None
    crawl_overview_file = os.path.join(
        dir_path, 'config', 'crawl_overview.yml')
    if os.path.isfile(crawl_overview_file):
        crawl_yml = readConfig(crawl_overview_file)
    else:
        print('No yml file')

    descript_lang = None
    descript_file = os.path.join(dir_path, 'config', 'crawl_overview_lang.yml')
    if os.path.isfile(descript_file):
        descript_lang = readConfig(descript_file)

    export_tabs = None
    export_tabs_file = os.path.join(dir_path, 'config', 'export_tabs.yml')
    if os.path.isfile(export_tabs_file):
        export_tabs = readConfig(export_tabs_file)

    domain = ""
    try:
        if len(sys.argv) > 1:
            domain = sys.argv[1]
    except KeyError as e:
        print('No sys arg domain given, continuing with domain list.')

    domains = {}
    domain_ints = []

    if domain != "":
        pass
    else:
        i = 1
        for domain in os.listdir(profiles):
            donts = ['.DS_Store', 'ORGS',
                     'ProjectInstanceData', 'crawl.seospiderconfig']
            if domain in donts:
                continue
            domains[i] = domain
            domain_ints.append(i)
            i += 1

        for k, v in domains.items():
            print(k, ": ", v)

        while True:
            domain_id = int(input('What domain? [int]'))
            if domain_id in domain_ints:
                try:
                    domain = domains[domain_id]
                except KeyError as e:
                    print('Wrong id!!')
                    sys.exit()

                print(f"Okay, processing {domain}")
                break

        profile = os.path.join(profiles, domain)
        config_file = os.path.join(profile, "config.yml")
        if os.path.isfile(config_file):
            c = readConfig(config_file)

            if c.config['active'] == 1:
                m = CreateWord(config=c.config, export_tabs=export_tabs.config,
                               bulk_export=bulk_export.config, crawl_yml=crawl_yml.config,
                               descript_lang=descript_lang.config)
            else:
                print("Skip: {}".format(profile))

    # profile = os.path.join(profiles, "oesterbaron.nl")
    # config_file = os.path.join(profile, "config.yml")
    # c = readConfig(config_file)
    # m = CreateWord(c.config)
    # m.dp.print_table_styles()
    # m.dp.print_paragraph_styles()
    # m.dp.print_list_styles()
